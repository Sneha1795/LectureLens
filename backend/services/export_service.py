from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os
import urllib.request
from fpdf import FPDF

FONT_DIR = "backend/fonts"
REGULAR_FONT_PATH = os.path.join(FONT_DIR, "DejaVuSans.ttf")
BOLD_FONT_PATH = os.path.join(FONT_DIR, "DejaVuSans-Bold.ttf")

def ensure_fonts():
    os.makedirs(FONT_DIR, exist_ok=True)
    if not os.path.exists(REGULAR_FONT_PATH):
        try:
            urllib.request.urlretrieve(
                "https://cdn.jsdelivr.net/gh/prawnpdf/prawn@master/data/fonts/DejaVuSans.ttf",
                REGULAR_FONT_PATH
            )
        except Exception as e:
            # Fallback if connection fails
            raise IOError(f"Failed to download DejaVuSans.ttf: {e}")

    if not os.path.exists(BOLD_FONT_PATH):
        try:
            urllib.request.urlretrieve(
                "https://cdn.jsdelivr.net/gh/prawnpdf/prawn@master/data/fonts/DejaVuSans-Bold.ttf",
                BOLD_FONT_PATH
            )
        except Exception:
            # Fallback: copy regular to bold if bold download fails
            import shutil
            shutil.copy(REGULAR_FONT_PATH, BOLD_FONT_PATH)

def generate_pdf(filename: str, full_text: str, keywords: list, summary: str, transcript: list) -> BytesIO:
    ensure_fonts()
    pdf = FPDF()
    pdf.add_page()

    # Load Unicode TrueType font to prevent encoding errors
    pdf.add_font("DejaVuSans", "", REGULAR_FONT_PATH)
    pdf.add_font("DejaVuSans", "B", BOLD_FONT_PATH)

    # Title
    pdf.set_font("DejaVuSans", "B", 16)
    pdf.cell(0, 10, "LectureLens Notes", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Filename
    pdf.set_font("DejaVuSans", "", 10)
    pdf.cell(0, 8, f"Lecture: {filename}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Summary Section
    if summary:
        pdf.set_font("DejaVuSans", "B", 12)
        pdf.cell(0, 10, "Summary & Notes", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("DejaVuSans", "", 10)
        pdf.multi_cell(0, 6, summary)
        pdf.ln(5)

    # Keywords Section
    if keywords:
        pdf.set_font("DejaVuSans", "B", 12)
        pdf.cell(0, 10, "Key Topics", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("DejaVuSans", "", 10)
        for kw in keywords:
            minutes = int(kw["timestamp"] // 60)
            seconds = int(kw["timestamp"] % 60)
            time_str = f"{minutes:02d}:{seconds:02d}"
            pdf.cell(0, 6, f"- {kw['keyword']}  [{time_str}]", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

    # Transcript Section
    if transcript:
        pdf.set_font("DejaVuSans", "B", 12)
        pdf.cell(0, 10, "Transcript", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("DejaVuSans", "", 10)
        for seg in transcript:
            minutes = int(seg["start"] // 60)
            seconds = int(seg["start"] % 60)
            time_str = f"{minutes:02d}:{seconds:02d}"
            pdf.multi_cell(0, 6, f"[{time_str}] {seg['text']}")

    pdf_bytes = pdf.output()
    return BytesIO(pdf_bytes)

def generate_docx(filename: str, full_text: str, keywords: list, summary: str, transcript: list) -> BytesIO:
    doc = Document()

    # Title
    title = doc.add_heading("LectureLens Notes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filename
    doc.add_paragraph(f"Lecture: {filename}")
    doc.add_paragraph("")

    # Summary Section
    if summary:
        doc.add_heading("Summary & Notes", level=1)
        doc.add_paragraph(summary)
        doc.add_paragraph("")

    # Keywords Section
    if keywords:
        doc.add_heading("Key Topics", level=1)
        for kw in keywords:
            minutes = int(kw["timestamp"] // 60)
            seconds = int(kw["timestamp"] % 60)
            time_str = f"{minutes:02d}:{seconds:02d}"
            doc.add_paragraph(f"• {kw['keyword']}  [{time_str}]")
        doc.add_paragraph("")

    # Transcript Section
    if transcript:
        doc.add_heading("Transcript", level=1)
        for seg in transcript:
            minutes = int(seg["start"] // 60)
            seconds = int(seg["start"] % 60)
            time_str = f"{minutes:02d}:{seconds:02d}"
            doc.add_paragraph(f"[{time_str}] {seg['text']}")
        doc.add_paragraph("")

    # Save document in memory to prevent local write and path traversal security vulnerabilities
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer